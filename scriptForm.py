from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx2pdf import convert
import pandas as pd
from pathlib import Path
import re

TEMPLATE = Path("plantilla.docx")
EXCEL = Path("datos.xlsx")       
HOJA = "Hoja1"
OUT_DOCX = Path("out_docx")
OUT_PDF = Path("out_pdf")

PLACEHOLDER_PATTERN = re.compile(r"\{\{(.*?)\}\}")  
def replace_in_paragraph(paragraph, mapping):
    if not paragraph.text:
        return
    text = paragraph.text
    for k, v in mapping.items():
        text = text.replace(f"{{{{{k}}}}}", str(v))
    for _ in range(len(paragraph.runs)):
        paragraph.runs[0].clear()
        paragraph.runs[0].element.getparent().remove(paragraph.runs[0].element)
    run = paragraph.add_run(text)

def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield parent.paragraphs[len(list(parent_elm.iterchildren())) - len(list(parent_elm.iterchildren()))]  

def replace_everywhere(doc: Document, mapping: dict):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)

def render_one(mapping: dict, out_docx_path: Path):
    doc = Document(TEMPLATE)
    replace_everywhere(doc, mapping)
    out_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx_path)

def main():
    df = pd.read_excel(EXCEL, sheet_name=HOJA, dtype=str).fillna("")
    OUT_DOCX.mkdir(exist_ok=True, parents=True)
    OUT_PDF.mkdir(exist_ok=True, parents=True)

    generated_paths = []
    for i, row in df.iterrows():
        mapping = {col: row[col] for col in df.columns}

        nombre_base = f"{i+1:03d}_{row.get('NOMBRE','doc')}".strip().replace(" ", "_")
        out_docx = OUT_DOCX / f"{nombre_base}.docx"
        out_pdf  = OUT_PDF  / f"{nombre_base}.pdf"

        render_one(mapping, out_docx)
        generated_paths.append((out_docx, out_pdf))

    for docx_path, pdf_path in generated_paths:
        convert(str(docx_path), str(pdf_path))

    print(f"Listo. DOCX en: {OUT_DOCX.resolve()}")
    print(f"PDF  en: {OUT_PDF.resolve()}")

if __name__ == "__main__":
    main()
